# -*- coding: utf-8 -*-

from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from mailmerge import MailMerge
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from django.shortcuts import render
from django.http import HttpResponse, JsonResponse, HttpResponseRedirect, FileResponse
import time
import datetime
import json
import os

import io

import math

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
UPLOAD_DIR = BASE_DIR + '/upload/'
# sys.path.append(BASE_DIR)


# from django.utils.http import urlquote


def data_input(request):
    if request.method == "GET":
        data = {}

        return render(request, 'data_input.html', data)


def data_input_com(request):
    if request.method == "GET":
        data = {}

        return render(request, 'data_input_com.html', data)


def get_test_data(request):
    if request.method == 'POST':
        Response = {}
        try:
            data = request.POST
            template_file_path = BASE_DIR + '/db_server/merge_file/template.docx'

            type_s = data.getlist("type")
            pipe_diameter_s = data.getlist("pipe_diameter")
            quantity_of_flow_s = data.getlist("quantity_of_flow")
            river_way_s = data.getlist("river_way")
            have_or_not_s = data.getlist("have_or_not")

            process_method = "□"
            not_process_method = "□"
            if data.get("process_method") == "self":
                process_method = "☑"
                not_process_method = "□"
            elif data.get("process_method") == "delegation":
                process_method = "□"
                not_process_method = "☑"

            process_ammount = "□"
            not_process_ammount = "□"
            if data.get("process_ammount") == "part":
                process_ammount = "☑"
                not_process_ammount = "□"
            elif data.get("process_ammount") == "all":
                process_ammount = "□"
                not_process_ammount = "☑"

            drainage = data.getlist("drainage")
            drainage1 = "□"
            drainage2 = "□"
            drainage3 = "□"
            drainage4 = "□"
            for i in drainage:
                if i == "drainage1":
                    drainage1 = "☑"
                elif i == "drainage2":
                    drainage2 = "☑"
                elif i == "drainage3":
                    drainage3 = "☑"
                elif i == "drainage4":
                    drainage4 = "☑"

            residence_area = data.get("residence_area")
            residence_area = residence_area if residence_area != "" and residence_area != "0" else "/"
            factory_area = data.get("factory_area")
            factory_area = factory_area if residence_area != "" and residence_area != "0" else "/"
            office_building_area = data.get("office_building_area")
            office_building_area = office_building_area if residence_area != "" and residence_area != "0" else "/"
            catering_area = data.get("catering_area")
            catering_area = catering_area if residence_area != "" and residence_area != "0" else "/"

            # 获取模版文件
            document = MailMerge(template_file_path)

            # 将获取的数据设置到模版的域中
            document.merge(
                applyer=data.get("applyer"),
                project=data.get("project"),
                date=data.get("date"),
                post_legal_representative=data.get(
                    "post_legal_representative"),
                cell_phone=data.get("cell_phone"),
                project_addr=data.get("project_addr"),
                zongshui=data.get("zongshui"),
                tap_water=data.get("tap_water"),
                self_owned=data.get("self_owned"),
                total_water=data.get("total_water"),
                total_catering=data.get("total_catering"),
                domastic_wastewater=data.get("domastic_wastewater"),
                process_method=process_method,
                not_process_method=not_process_method,
                process_ammount=process_ammount,
                not_process_ammount=not_process_ammount,
                process_technology=data.get("process_technology"),
                site_area=data.get("site_area"),
                overall_floorage=data.get("overall_floorage"),
                residence_area=residence_area,
                factory_area=factory_area,
                office_building_area=office_building_area,
                catering_area=catering_area,
                main_product=data.get("main_product"),
                raw_material=data.get("raw_material"),
                construction_period=data.get("construction_period"),
                completion_date=data.get("completion_date"),

                type_1=type_s[0],
                pipe_diameter_1=pipe_diameter_s[0],
                quantity_of_flow_1=quantity_of_flow_s[0],
                river_way_1=river_way_s[0],
                have_or_not_1=have_or_not_s[0],

                type_2=type_s[1],
                pipe_diameter_2=pipe_diameter_s[1],
                quantity_of_flow_2=quantity_of_flow_s[1],
                river_way_2=river_way_s[1],
                have_or_not_2=have_or_not_s[1],

                type_3=type_s[2],
                pipe_diameter_3=pipe_diameter_s[2],
                quantity_of_flow_3=quantity_of_flow_s[2],
                river_way_3=river_way_s[2],
                have_or_not_3=have_or_not_s[2],

                type_4=type_s[3],
                pipe_diameter_4=pipe_diameter_s[3],
                quantity_of_flow_4=quantity_of_flow_s[3],
                river_way_4=river_way_s[3],
                have_or_not_4=have_or_not_s[3],

                type_5=type_s[4],
                pipe_diameter_5=pipe_diameter_s[4],
                quantity_of_flow_5=quantity_of_flow_s[4],
                river_way_5=river_way_s[4],
                have_or_not_5=have_or_not_s[4],

                type_6=type_s[5],
                pipe_diameter_6=pipe_diameter_s[5],
                quantity_of_flow_6=quantity_of_flow_s[5],
                river_way_6=river_way_s[5],
                have_or_not_6=have_or_not_s[5],

                type_7=type_s[6],
                pipe_diameter_7=pipe_diameter_s[6],
                quantity_of_flow_7=quantity_of_flow_s[6],
                river_way_7=river_way_s[6],
                have_or_not_7=have_or_not_s[6],

                type_8=type_s[7],
                pipe_diameter_8=pipe_diameter_s[7],
                quantity_of_flow_8=quantity_of_flow_s[7],
                river_way_8=river_way_s[7],
                have_or_not_8=have_or_not_s[7],

                levels=data.get("levels"),
                drainage1=drainage1,
                drainage2=drainage2,
                drainage3=drainage3,
                drainage4=drainage4,

            )
            pics_description = data.getlist("pics_description")

            time_text = time.strftime("%Y%m%d%H%M%S", time.localtime())
            result_file_path = BASE_DIR + '/db_server/merge_file/temp/' + time_text + '.docx'

            document.write(result_file_path)

            # 处理图片
            doc = Document(result_file_path)
            tables = doc.tables
            pic_table = tables[3]

            file_data = request.FILES.getlist('file')

            for index, fl in enumerate(file_data):
                filename = fl._get_name()

                if index > 0:
                    pic_table.add_row()
                    pic_table.add_row()

                cell = pic_table.cell(index * 2, 0)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell.paragraphs[0].add_run()
                picture = run.add_picture(fl)
                picture.height = Cm(6)
                picture.width = Cm(8)

                cell = pic_table.cell(index * 2 + 1, 0)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                run = cell.paragraphs[0].add_run()
                if index < len(pics_description):
                    text = run.add_text(pics_description[index])

            doc.save(result_file_path)

        except Exception as e:
            err_message = str(e)

            # 代码 0 其他错误，并把错误信息返回
            Response['code'] = 0
            Response['msg'] = err_message

            return HttpResponse(json.dumps(Response))

        else:   # 代码无异常
            Response['code'] = 14000
            Response['msg'] = "成功"
            Response['url'] = "/db_server/download/" + time_text

            return HttpResponse(json.dumps(Response))


def download(request, id):
    if request.method == "GET":
        response = b''
        file_name = BASE_DIR + '/db_server/merge_file/temp/' + id + '.docx'
        with open(file_name, 'rb') as file:
            response = HttpResponse(
                file, content_type='application/octet-stream')
            # 文件流类型
            filename = id + ".docx"
            response['Content-Disposition'] = 'attachment; filename={}'.format(
                filename)
            # 表示带有文件附件，指定了文件名；浏览器会自动下载

        # os.remove(file_name)
        return response
